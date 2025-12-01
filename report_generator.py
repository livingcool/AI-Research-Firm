from fpdf import FPDF
from docx import Document
import os

class PDFReport(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 15)
        self.cell(0, 10, 'Autonomous Research Report', 0, 1, 'C')
        self.ln(10)

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

def generate_pdf(content: str, filename: str = "report.pdf"):
    """
    Generates a PDF report from the markdown content.
    Note: Basic markdown parsing.
    """
    pdf = PDFReport()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    # Simple line-by-line write (Enhancement: Use a proper MD to PDF library)
    for line in content.split('\n'):
        # Handle headers roughly
        if line.startswith('#'):
            pdf.set_font("Arial", 'B', 14)
            pdf.cell(0, 10, line.replace('#', '').strip(), 0, 1)
            pdf.set_font("Arial", size=12)
        else:
            # Handle long lines
            pdf.multi_cell(0, 10, line)
            
    output_path = os.path.join(tempfile.gettempdir(), filename)
    pdf.output(output_path)
    return output_path

def generate_docx(content: str, filename: str = "report.docx"):
    """
    Generates a Word document from the markdown content.
    """
    doc = Document()
    doc.add_heading('Autonomous Research Report', 0)
    
    for line in content.split('\n'):
        if line.startswith('#'):
            level = min(line.count('#'), 9)
            doc.add_heading(line.replace('#', '').strip(), level=level)
        else:
            doc.add_paragraph(line)
            
    output_path = os.path.join(tempfile.gettempdir(), filename)
    doc.save(output_path)
    return output_path

import tempfile
